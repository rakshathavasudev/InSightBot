import os
import time
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document
from langchain_pinecone import PineconeVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
import pinecone
from dotenv import load_dotenv
from pinecone import Pinecone
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.schema import Document as LangchainDocument
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# Load environment variables
load_dotenv()
# os.environ['PINECONE_API_KEY'] = os.environ.get('PINECONE_API_KEY')

class DocumentProcessor:
    def __init__(self, index_name: str, namespace: str, chunk_size: int = 1000, chunk_overlap: int = 30):
        self.index_name = index_name
        self.namespace = namespace
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def delete_files_in_directory(self, directory_path: str):
        """
        Delete all files in a given directory.
        """
        files = os.listdir(directory_path)
        print('Files in dir:', files)
        for file in files:
            file_path = os.path.join(directory_path, file)
            if os.path.isfile(file_path):
                os.remove(file_path)
        print("All files deleted successfully.")

    def process_directory(self, directory_path: str):
        data = []
        
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                print(f"Processing file: {file_path}")
                
                # Skip empty files
                if os.path.getsize(file_path) == 0:
                    print(f"Warning: The file {file_path} is empty and will be skipped.")
                    continue
                
                try:
                    # Process .txt files
                    if file.endswith(".txt"):
                        with open(file_path, "r", encoding="utf-8") as f:
                            file_data = f.read()
                    
                    # Process .pdf files
                    elif file.endswith(".pdf"):
                        reader = PdfReader(file_path)
                        file_data = ""
                        for page in reader.pages:
                            # extract_text() may return None for scanned or empty pages
                            page_text = page.extract_text() or ""
                            file_data += page_text
                        # If no text extracted and OCR is available, try OCR as a fallback
                        if not file_data.strip() and OCR_AVAILABLE:
                            try:
                                print(f"No text found in PDF via PdfReader, attempting OCR for {file_path}")
                                images = convert_from_path(file_path)
                                ocr_text = []
                                for img in images:
                                    text = pytesseract.image_to_string(img)
                                    ocr_text.append(text)
                                ocr_result = "\n".join(ocr_text)
                                if ocr_result.strip():
                                    print(f"OCR succeeded for {file_path}")
                                    file_data = ocr_result
                                else:
                                    print(f"OCR returned no text for {file_path}")
                            except Exception as ocr_e:
                                print(f"OCR failed for {file_path}: {ocr_e}")
                    
                    # Process .docx files
                    elif file.endswith(".docx"):
                        document = DocxDocument(file_path)
                        file_data = "\n".join([para.text for para in document.paragraphs])
                    
                    else:
                        print(f"Unsupported file format: {file}")
                        continue
                    
                    # Ensure valid content was extracted
                    if not file_data.strip():
                        print(f"Warning: No content extracted from {file_path}. Skipping.")
                        continue
                    
                    # Append the processed data
                    data.append({"File": file_path, "Data": file_data})
                
                except Exception as e:
                    print(f"Error processing file {file_path}: {e}")
        return data

    def prepare_data(self, documents):
        # Prepare the text for embedding
        document_data = []
        for document in documents:
            # Ensure 'Data' exists and is non-empty
            if 'Data' not in document or not document['Data']:
                print(f"Skipping document due to missing or empty 'Data': {document}")
                continue
            
            # Ensure the first element in 'Data' is a Document object
            if not isinstance(document['Data'][0], str):
                print(f"Skipping document due to invalid data type in 'Data': {document}")
                continue

            # Extract metadata and content
            document_source = document['File']
            document_content = document['Data']

            # Normalize path so this works on Windows and Unix
            norm_source = os.path.normpath(document_source)
            parts = norm_source.split(os.path.sep)
            file_name = parts[-1]
            folder_names = parts[1:-1] if len(parts) > 2 else []

            doc = LangchainDocument(
                page_content=f"<Source>\n{document_source}\n</Source>\n\n<Content>\n{document_content}\n</Content>",
                metadata={
                    "file_name": file_name,
                    "parent_folder": folder_names[-1] if folder_names else "",
                    "folder_names": folder_names
                }
            )
            document_data.append(doc)

        return document_data

    def chunk_data(self, docs):
        """
        Split documents into chunks based on the provided chunk size and overlap.
        """
        textsplitter = RecursiveCharacterTextSplitter(chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
        docs = textsplitter.split_documents(docs)
        return docs

    def upsert_vectorstore_to_pinecone(self, document_data, embeddings, index_name, namespace):
        # # Initialize Pinecone connection with the new API structure
        pc = pinecone.Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))

        # Check if the namespace exists in the index
        index = pc.Index(self.index_name)

        # Check if the namespace exists by listing the namespaces (or by trying to query)
        namespaces = index.describe_index_stats().get('namespaces', [])
        max_retries = 5
        wait_time = 2000
        if self.namespace in namespaces:
            print(f"Namespace '{self.namespace}' found. Deleting vector data...")
            index.delete(namespace=self.namespace, delete_all=True)  # Initiates deletion

            # Polling to ensure deletion completes
            for attempt in range(max_retries):
                namespaces = index.describe_index_stats().get('namespaces', [])
                if self.namespace not in namespaces:
                    print(f"Namespace '{self.namespace}' deletion confirmed.")
                    break
                time.sleep(wait_time)  # Wait before re-checking
            else:
                raise RuntimeError(f"Failed to delete namespace '{self.namespace}' after {max_retries} retries.")

        else:
            print(f"Namespace '{self.namespace}' does not exist. Proceeding with upsert.")

        # Create or replace the vector store
        vectorstore_from_documents = PineconeVectorStore.from_documents(
            document_data,
            embeddings,
            index_name=self.index_name,
            namespace=self.namespace
        )
        print(f"Vector store type: {type(vectorstore_from_documents)}")

        # Optionally, return the vector store if needed
        return vectorstore_from_documents

    def initialize_pinecone(self):
        """
        Initialize Pinecone with the provided index name.
        """
        # Initialize Pinecone instance
        pc = Pinecone(api_key=os.environ.get('PINECONE_API_KEY'))

        # Check if the index exists
        if self.index_name not in [index.name for index in pc.list_indexes().indexes]:
            raise ValueError(f"Index '{self.index_name}' does not exist. Please create it first.")

        # Connect to the specified index
        pinecone_index = pc.Index(self.index_name)
        return pinecone_index
